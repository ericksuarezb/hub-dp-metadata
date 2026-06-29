from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.models import DocumentArtifact, ProcessStep, StepDocumentModel, StructuralDocumentModel

TITLE_STYLE = "CustomTitle"
SECTION_STYLE = "H2"
SMALL_STYLE = "Small"
BODY_STYLE = "Normal"
TABLE_STYLE = "Table Grid"
NAVY = "162B44"
BLUE = "2B6CB0"
PURPLE = "6B21A8"
PURPLE_BG = "F3E8FF"
GREEN = "15803D"
ORANGE = "C2410C"
GRAY = "505050"
LIGHT_GRAY = "F2F4F6"
ALT_ROW = "EBF2FA"
BORDER = "D4D4D4"
TOTAL_TABLE_WIDTH = Inches(9.5)
STEP_COLORS = {
    1: ("DBEAFE", "1E40AF"),
    2: ("E9D5FF", "6B21A8"),
    3: ("FEF3C7", "92400E"),
    4: ("D1FAE5", "065F46"),
    5: ("FCE7F3", "9D174D"),
    6: ("E0E7FF", "3730A3"),
    7: ("FEE2E2", "991B1B"),
}

STEP_SECTION_TITLES = [
    "1. Fuentes de datos",
    "2. Flujo general del proceso",
    "3. Matriz de trazabilidad",
    "4. Transformaciones",
    "5. Reglas de negocio",
    "6. Proceso paso a paso detallado",
    "7. Criterio de aceptacion",
]

STRUCTURAL_SECTION_TITLES = [
    "1. Referencia rapida y ficha del producto",
    "2. Fuentes de datos",
    "3. Flujo general del proceso",
    "4. Matriz de trazabilidad",
    "5. Transformaciones",
    "6. Reglas de negocio",
    "7. Proceso paso a paso detallado",
    "8. Convenciones y navegacion",
]


def render_structural_docx(
    model: StructuralDocumentModel,
    template_path: str | Path,
    output_path: str | Path,
) -> DocumentArtifact:
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"No existe la plantilla estructural: {template}")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template))
    _clear_document(doc)
    _configure_document_styles(doc)
    _configure_sections(doc, model.title)

    _add_structural_cover(doc, model)
    _add_structural_section_1(doc, model)
    _add_structural_sources(doc, model)
    _add_structural_flow(doc, model)
    _add_structural_traceability(doc, model)
    _add_structural_transformations(doc, model)
    _add_structural_rules(doc, model)
    _add_process_section(doc, model.global_steps)
    _add_structural_navigation(doc, model)

    doc.save(str(output))
    return DocumentArtifact(
        path=str(output),
        document_kind="structural",
        section_titles=STRUCTURAL_SECTION_TITLES,
        section4_fields=[item.field_name for item in model.final_transformations],
        section5_fields=[item.field_name for item in model.final_transformations],
        referenced_rule_ids=sorted({item.rule_id for item in model.final_transformations if item.rule_id}),
        document_text="\n".join(_iter_document_text(doc)),
    )


def render_step_docx(
    model: StepDocumentModel,
    template_path: str | Path,
    output_path: str | Path,
) -> DocumentArtifact:
    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"No existe la plantilla STEP: {template}")

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(template))
    _clear_document(doc)
    _configure_document_styles(doc)
    _configure_sections(doc, f"STEP - {model.module_name}")

    _add_cover(doc, model)
    _add_sources_section(doc, model)
    _add_flow_section(doc, model)
    _add_traceability_section(doc, model)
    _add_transformations_section(doc, model)
    _add_rules_section(doc, model)
    _add_process_section(doc, model.process_steps)
    _add_acceptance_section(doc, model)

    doc.save(str(output))
    return DocumentArtifact(
        path=str(output),
        document_kind="step",
        section_titles=STEP_SECTION_TITLES,
        section4_fields=[item.field_name for item in model.transformations],
        section5_fields=[item.field_name for item in model.transformations],
        referenced_rule_ids=sorted({item.rule_id for item in model.transformations if item.rule_id}),
        document_text="\n".join(_iter_document_text(doc)),
    )


def _add_structural_cover(doc: Document, model: StructuralDocumentModel) -> None:
    title = doc.add_paragraph()
    _apply_paragraph_style(title, TITLE_STYLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "ESPECIFICACION FUNCIONAL", bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    _apply_paragraph_style(subtitle, BODY_STYLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, model.product_name, bold=True, color=BLUE)

    intro = doc.add_paragraph()
    _apply_paragraph_style(intro, BODY_STYLE)
    _add_run(
        intro,
        f"{model.title}. Este documento resume la vista estructural del producto y la secuencia de modulos que lo construyen.",
    )

    toc = doc.add_paragraph()
    _apply_paragraph_style(toc, BODY_STYLE)
    _add_run(toc, "Contenido:", bold=True, color=NAVY)
    for section in STRUCTURAL_SECTION_TITLES:
        paragraph = doc.add_paragraph()
        _apply_paragraph_style(paragraph, SMALL_STYLE)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        _add_run(paragraph, section)

    _add_layer_identification_box(doc, model)
    doc.add_page_break()


def _add_structural_section_1(doc: Document, model: StructuralDocumentModel) -> None:
    _add_h2(doc, STRUCTURAL_SECTION_TITLES[0])
    block = doc.add_paragraph()
    _apply_paragraph_style(block, BODY_STYLE)
    _add_run(block, "FICHA DEL PRODUCTO", bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=2)
    _style_table(table)
    _header_row(table.rows[0], ["Campo", "Contenido"], fill=NAVY)
    for key, value in model.product_sheet.items():
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
    _finalize_table(table)

    block = doc.add_paragraph()
    _apply_paragraph_style(block, BODY_STYLE)
    _add_run(block, "TABLAS CRYSTAL DE SALIDA", bold=True, color=PURPLE)
    crystal = doc.add_table(rows=1, cols=2)
    _style_table(crystal)
    _header_row(crystal.rows[0], ["Tabla Crystal", "Descripcion"], fill=PURPLE)
    for item in model.output_tables:
        row = crystal.add_row().cells
        row[0].text = item.tabla
        row[1].text = item.descripcion
        for cell in row:
            _shade_cell(cell, PURPLE_BG)
    _finalize_table(crystal)

    block = doc.add_paragraph()
    _apply_paragraph_style(block, BODY_STYLE)
    _add_run(block, "REFERENCIA RAPIDA", bold=True, color=BLUE)
    quick = doc.add_table(rows=1, cols=2)
    _style_table(quick)
    _header_row(quick.rows[0], ["Pregunta", "Respuesta"], fill=BLUE)
    for key, value in model.quick_reference.items():
        row = quick.add_row().cells
        row[0].text = key
        row[1].text = value
    _finalize_table(quick)


def _add_layer_identification_box(doc: Document, model: StructuralDocumentModel) -> None:
    box = doc.add_table(rows=1, cols=1)
    _style_table(box)
    _shade_cell(box.rows[0].cells[0], PURPLE_BG)
    title_paragraph = box.rows[0].cells[0].paragraphs[0]
    _apply_paragraph_style(title_paragraph, BODY_STYLE)
    _add_run(title_paragraph, "Identificacion de capas de datos", bold=True, color=PURPLE)

    for note in model.layer_identification_notes:
        paragraph = box.rows[0].cells[0].add_paragraph()
        _apply_paragraph_style(paragraph, BODY_STYLE)
        _add_run(paragraph, f"- {note}")

    if model.mapped_variables:
        variables_title = box.rows[0].cells[0].add_paragraph()
        _apply_paragraph_style(variables_title, BODY_STYLE)
        _add_run(variables_title, "Variables mapeadas", bold=True, color=NAVY)
        for item in model.mapped_variables:
            paragraph = box.rows[0].cells[0].add_paragraph()
            _apply_paragraph_style(paragraph, BODY_STYLE)
            _add_run(paragraph, f"- {item}", color=GRAY)


def _add_structural_sources(doc: Document, model: StructuralDocumentModel) -> None:
    _add_sources_table(doc, STRUCTURAL_SECTION_TITLES[1], model.global_sources)


def _add_structural_flow(doc: Document, model: StructuralDocumentModel) -> None:
    _add_h2(doc, STRUCTURAL_SECTION_TITLES[2])
    band = doc.add_paragraph()
    _apply_paragraph_style(band, BODY_STYLE)
    _add_run(band, "MAPA DE MODULOS Y SECUENCIA FUNCIONAL DEL PIPELINE", bold=True, color="FFFFFF")
    if band.runs:
        _shade_paragraph(band, NAVY)

    subtitle = doc.add_paragraph()
    _apply_paragraph_style(subtitle, BODY_STYLE)
    _add_run(
        subtitle,
        "Vista resumida de los modulos logicos del producto, su intencion funcional y su papel documental dentro del pipeline.",
    )

    timeline = doc.add_paragraph()
    _apply_paragraph_style(timeline, SMALL_STYLE)
    timeline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    timeline_parts = []
    for index, module in enumerate(model.module_sequence, start=1):
        marker = "Principal" if module.document_status == "Principal" else "STEP"
        timeline_parts.append(f"P{index} {module.module_id} [{marker}]")
    _add_run(timeline, "  ->  ".join(timeline_parts), bold=True, color=BLUE)

    modules = doc.add_table(rows=1, cols=6)
    _style_table(modules)
    _header_row(modules.rows[0], ["Modulo", "Intencion", "Dependencias", "SQL", "Salida", "Estado"])
    for module in model.module_sequence:
        row = modules.add_row().cells
        row[0].text = f"{module.module_id}\n{module.module_name}"
        row[1].text = module.module_intention
        row[2].text = ", ".join(module.depends_on) if module.depends_on else "—"
        row[3].text = module.sql_path
        row[4].text = ", ".join(module.output_tables) if module.output_tables else "—"
        row[5].text = f"{module.document_role} / {module.document_status}"
        if module.document_status == "Principal":
            _shade_cell(row[5], "D1FAE5")
    _finalize_table(modules)

    _add_flow_table(doc, model.global_steps, section_title=None)


def _add_structural_traceability(doc: Document, model: StructuralDocumentModel) -> None:
    _add_h2(doc, STRUCTURAL_SECTION_TITLES[3])
    table = doc.add_table(rows=1, cols=8)
    _style_table(table)
    _header_row(table.rows[0], ["Campo", "P1", "P2", "P3", "P4", "P5", "P6", "P7"])
    for item in model.final_transformations:
        row = table.add_row().cells
        row[0].text = item.field_name
        for offset in range(1, 8):
            row[offset].text = "●" if f"Paso {offset}" in item.participates_in_steps else ""
    _finalize_traceability_table(table, step_count=7)


def _add_structural_transformations(doc: Document, model: StructuralDocumentModel) -> None:
    _add_transformations_table(doc, STRUCTURAL_SECTION_TITLES[4], model.final_transformations)


def _add_structural_rules(doc: Document, model: StructuralDocumentModel) -> None:
    _add_rules_table(doc, STRUCTURAL_SECTION_TITLES[5], model.global_rules, "No se detectaron reglas explicitas a nivel global.")


def _add_structural_navigation(doc: Document, model: StructuralDocumentModel) -> None:
    _add_h2(doc, STRUCTURAL_SECTION_TITLES[7])
    guide = doc.add_table(rows=1, cols=2)
    _style_table(guide)
    _header_row(guide.rows[0], ["Necesidad", "Ir a"])
    for note in model.navigation_notes:
        left, right = [part.strip() for part in note.split("->", 1)]
        row = guide.add_row().cells
        row[0].text = left
        row[1].text = right
    _finalize_table(guide)


def _clear_document(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _configure_document_styles(doc: Document) -> None:
    _set_style(doc, "Normal", 17, GRAY, False)
    _set_style(doc, "CustomTitle", 26, NAVY, True)
    _set_style(doc, "H2", 21, BLUE, True)
    _set_style(doc, "Small", 17, GRAY, False)


def _configure_sections(doc: Document, header_text: str) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        _configure_header(section, header_text)
        _configure_footer(section)


def _add_cover(doc: Document, model: StepDocumentModel) -> None:
    title = doc.add_paragraph()
    _apply_paragraph_style(title, TITLE_STYLE)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(title, "ESPECIFICACION FUNCIONAL", bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    _apply_paragraph_style(subtitle, BODY_STYLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(subtitle, model.module_name, bold=True, color=BLUE)

    intro = doc.add_paragraph()
    _apply_paragraph_style(intro, BODY_STYLE)
    _add_run(
        intro,
        f"Modulo {model.module_id} del producto {model.product_name}. Intencion: {model.module_intention}",
    )
    if model.depends_on:
        depends = doc.add_paragraph()
        _apply_paragraph_style(depends, BODY_STYLE)
        _add_run(depends, f"Depende de: {', '.join(model.depends_on)}")
    if model.output_tables:
        outputs = doc.add_paragraph()
        _apply_paragraph_style(outputs, BODY_STYLE)
        _add_run(outputs, f"Salida del modulo: {', '.join(model.output_tables)}", color=PURPLE)
    if model.publications:
        publication_summary = doc.add_paragraph()
        _apply_paragraph_style(publication_summary, SMALL_STYLE)
        _add_run(
            publication_summary,
            "Publicaciones detectadas en el archivo SQL: "
            + "; ".join(
                f"P{item.sequence}: {item.target_table} ({item.role})" for item in model.publications
            ),
            color=NAVY,
        )

    toc = doc.add_paragraph()
    _apply_paragraph_style(toc, BODY_STYLE)
    _add_run(toc, "Contenido:", bold=True, color=NAVY)
    for section in STEP_SECTION_TITLES:
        paragraph = doc.add_paragraph()
        _apply_paragraph_style(paragraph, SMALL_STYLE)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        _add_run(paragraph, section)
    doc.add_page_break()


def _add_sources_section(doc: Document, model: StepDocumentModel) -> None:
    _add_sources_table(doc, STEP_SECTION_TITLES[0], model.sources)


def _add_flow_section(doc: Document, model: StepDocumentModel) -> None:
    _add_h2(doc, STEP_SECTION_TITLES[1])
    if model.publications:
        _add_publications_table(doc, model.publications)
    _add_flow_table(doc, model.process_steps, section_title=None)


def _add_traceability_section(doc: Document, model: StepDocumentModel) -> None:
    _add_h2(doc, STEP_SECTION_TITLES[2])
    table = doc.add_table(rows=1, cols=7)
    _style_table(table)
    _header_row(table.rows[0], ["Campo", "P1", "P2", "P3", "P4", "P5", "P6"])
    for item in model.transformations:
        row = table.add_row().cells
        row[0].text = item.field_name
        for offset in range(1, 7):
            row[offset].text = "●" if f"Paso {offset}" in item.participates_in_steps else ""
    _finalize_traceability_table(table, step_count=6)


def _add_transformations_section(doc: Document, model: StepDocumentModel) -> None:
    _add_transformations_table(doc, STEP_SECTION_TITLES[3], model.transformations)


def _add_rules_section(doc: Document, model: StepDocumentModel) -> None:
    _add_rules_table(doc, STEP_SECTION_TITLES[4], model.rules, "No se detectaron reglas explicitas para este modulo.")


def _add_process_section(doc: Document, steps: Sequence[ProcessStep]) -> None:
    _add_process_cards(doc, STEP_SECTION_TITLES[5], steps)


def _add_acceptance_section(doc: Document, model: StepDocumentModel) -> None:
    _add_h2(doc, STEP_SECTION_TITLES[6])
    table = doc.add_table(rows=1, cols=2)
    _style_table(table)
    _header_row(table.rows[0], ["#", "Criterio"])
    for index, criterion in enumerate(model.acceptance_criteria, start=1):
        row = table.add_row().cells
        row[0].text = str(index)
        row[1].text = criterion
    _finalize_table(table)


def _add_publications_table(doc: Document, publications) -> None:
    block = doc.add_paragraph()
    _apply_paragraph_style(block, BODY_STYLE)
    _add_run(block, "PUBLICACIONES DEL MODULO", bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=4)
    _style_table(table)
    _header_row(table.rows[0], ["Orden", "Tipo", "Tabla publicada", "Estado"])
    for publication in publications:
        row = table.add_row().cells
        row[0].text = f"P{publication.sequence}"
        row[1].text = _publication_type_label(publication.statement_type)
        row[2].text = publication.target_table
        status = "Final" if publication.role == "final" else "Intermedia"
        stats = " + estadisticas recalculadas" if publication.has_compute_stats else ""
        row[3].text = f"{status}{stats}"
    _finalize_table(table)


def _add_sources_table(doc: Document, section_title: str, sources) -> None:
    _add_h2(doc, section_title)
    table = doc.add_table(rows=1, cols=7)
    _style_table(table)
    _header_row(
        table.rows[0],
        ["Alias", "Tabla insumo", "Capa", "Campos finales que genera", "Que contiene", "Se usa en", "Tabla Crystal destino"],
    )
    for source in sources:
        row = table.add_row().cells
        row[0].text = source.alias
        row[1].text = source.table_name
        row[2].text = source.layer
        row[3].text = ", ".join(source.fields_generated)
        row[4].text = source.contains_description
        row[5].text = ", ".join(source.used_in_steps)
        row[6].text = source.destination_table
    _finalize_table(table)


def _add_flow_table(doc: Document, steps: Sequence[ProcessStep], section_title: str | None = STRUCTURAL_SECTION_TITLES[2]) -> None:
    if section_title:
        _add_h2(doc, section_title)
    table = doc.add_table(rows=1, cols=5)
    _style_table(table)
    _header_row(table.rows[0], ["Paso", "Accion", "Depende de", "Tablas nuevas", "Produce"])
    for step in steps:
        row = table.add_row().cells
        row[0].text = str(step.number)
        row[1].text = step.title
        row[2].text = step.depends_on
        row[3].text = ", ".join(step.tables_involved[:4])
        row[4].text = step.result
    _finalize_step_flow_table(table)


def _add_transformations_table(doc: Document, section_title: str, transformations) -> None:
    _add_h2(doc, section_title)
    table = doc.add_table(rows=1, cols=9)
    _style_table(table)
    _header_row(table.rows[0], ["#", "Campo destino", "Tipo", "Subtipo", "Origen", "Campo origen", "Que se le hace", "Paso", "Regla"])
    for item in transformations:
        row = table.add_row().cells
        row[0].text = str(item.index)
        row[1].text = item.field_name
        row[2].text = item.field_type
        row[3].text = item.subtype
        row[4].text = item.origin
        row[5].text = " / ".join(item.source_fields) if item.source_fields else "—"
        row[6].text = item.description
        row[7].text = item.step
        row[8].text = item.rule_id or "—"
    _finalize_transformations_table(table)


def _add_rules_table(doc: Document, section_title: str, rules, empty_message: str) -> None:
    _add_h2(doc, section_title)
    table = doc.add_table(rows=1, cols=3)
    _style_table(table)
    _header_row(table.rows[0], ["ID", "Regla de negocio", "Se aplica en"])
    if rules:
        for rule in rules:
            row = table.add_row().cells
            row[0].text = rule.id
            row[1].text = rule.description
            row[2].text = rule.applies_in
    else:
        row = table.add_row().cells
        row[0].text = "—"
        row[1].text = empty_message
        row[2].text = "—"
    _finalize_table(table)


def _add_process_cards(doc: Document, section_title: str, steps: Sequence[ProcessStep]) -> None:
    _add_h2(doc, section_title)
    for index, step in enumerate(steps, start=1):
        card = doc.add_table(rows=10, cols=2)
        _style_table(card)
        labels = [
            "Encabezado",
            "Depende de",
            "Tablas involucradas",
            "Criterio de cruce",
            "Tipo de union",
            "Que significa",
            "Criterio de seleccion",
            "Campos que se extraen",
            "Reglas aplicadas",
            "Resultado",
        ]
        values = [
            f"PASO {step.number} - {step.title}. Objetivo: {step.objective}",
            step.depends_on,
            ", ".join(step.tables_involved),
            " | ".join(step.join_criteria) if step.join_criteria else "No aplica",
            " | ".join(step.join_type) if step.join_type else "No aplica",
            " | ".join(step.meaning) if step.meaning else "No aplica",
            " | ".join(step.selection_criteria) if step.selection_criteria else "No aplica",
            " | ".join(step.extracted_fields) if step.extracted_fields else "Ver Seccion 4",
            ", ".join(step.rule_ids) if step.rule_ids else "—",
            step.result,
        ]
        for row_index, (label, value) in enumerate(zip(labels, values)):
            card.rows[row_index].cells[0].text = label
            card.rows[row_index].cells[1].text = value
        _finalize_process_card(card)
        if index != len(steps):
            arrow = doc.add_paragraph()
            _apply_paragraph_style(arrow, BODY_STYLE)
            arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(arrow, "↓", size=18, bold=True, color=BLUE)


def _add_h2(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    _apply_paragraph_style(paragraph, SECTION_STYLE)
    _add_run(paragraph, text, bold=True, color=BLUE)


def _add_run(paragraph, text: str, size: int | None = None, bold: bool = False, color: str | None = "505050") -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_table(table) -> None:
    try:
        table.style = TABLE_STYLE
    except KeyError:
        table.style = "Table Grid"
    table.autofit = False
    _set_table_borders(table)
    _set_table_widths(table)


def _header_row(row, values, fill: str = "162B44") -> None:
    for cell, value in zip(row.cells, values):
        cell.text = value
        _shade_cell(cell, fill)
        if cell.paragraphs:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for child in list(tc_pr):
        if child.tag == qn("w:shd"):
            tc_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)
    _set_cell_border(cell)


def _shade_paragraph(paragraph, fill: str) -> None:
    paragraph_pr = paragraph._p.get_or_add_pPr()
    for child in list(paragraph_pr):
        if child.tag == qn("w:shd"):
            paragraph_pr.remove(child)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    paragraph_pr.append(shd)


def _publication_type_label(statement_type: str) -> str:
    if statement_type == "insert_overwrite":
        return "Publicacion por reemplazo"
    if statement_type == "create_table_as_select":
        return "Creacion con datos"
    return "Publicacion documentada"


def _apply_paragraph_style(paragraph, style_name: str) -> None:
    try:
        paragraph.style = style_name
    except KeyError:
        return


def _set_style(doc: Document, style_name: str, size_pt: int, color: str, bold: bool) -> None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return
    style.font.name = "Arial"
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def _configure_header(section, header_text: str) -> None:
    paragraph = section.header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(header_text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run.bold = True
    _set_paragraph_bottom_border(paragraph, BORDER)


def _configure_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    _append_page_number_field(run)


def _append_page_number_field(run) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _set_paragraph_bottom_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), BORDER)


def _set_table_widths(table) -> None:
    column_width = int(TOTAL_TABLE_WIDTH / max(len(table.columns), 1))
    for row in table.rows:
        for cell in row.cells:
            cell.width = column_width


def _set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), BORDER)


def _finalize_table(table, header_rows: int = 1, font_size: int = 12) -> None:
    for row_index, row in enumerate(table.rows):
        if row_index >= header_rows:
            fill = ALT_ROW if (row_index - header_rows) % 2 == 0 else "FFFFFF"
            for cell in row.cells:
                _shade_cell(cell, fill)
                for paragraph in cell.paragraphs:
                    _apply_paragraph_style(paragraph, BODY_STYLE)
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(font_size)
                        run.font.color.rgb = RGBColor.from_string(GRAY)


def _finalize_traceability_table(table, step_count: int) -> None:
    _finalize_table(table)
    for index in range(1, min(step_count, len(table.rows[0].cells) - 1) + 1):
        fill, text = STEP_COLORS[index]
        cell = table.rows[0].cells[index]
        _shade_cell(cell, fill)
        if cell.paragraphs:
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(text)
                run.font.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(12)
    for row in table.rows[1:]:
        for index in range(1, min(step_count, len(row.cells) - 1) + 1):
            if row.cells[index].text.strip() == "●":
                fill, text = STEP_COLORS[index]
                _shade_cell(row.cells[index], fill)
                for run in row.cells[index].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor.from_string(text)
                    run.font.bold = True


def _finalize_transformations_table(table) -> None:
    _finalize_table(table)
    for row in table.rows[1:]:
        type_value = row.cells[2].text.strip()
        if type_value == "D":
            _shade_cell(row.cells[2], "D1FAE5")
            color = GREEN
        elif type_value == "T":
            _shade_cell(row.cells[2], "FED7AA")
            color = ORANGE
        else:
            continue
        for run in row.cells[2].paragraphs[0].runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.bold = True
        step_text = row.cells[7].text.strip()
        step_number = _extract_step_number(step_text)
        if step_number in STEP_COLORS:
            fill, text = STEP_COLORS[step_number]
            _shade_cell(row.cells[7], fill)
            for run in row.cells[7].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(text)
                run.font.bold = True


def _finalize_step_flow_table(table) -> None:
    _finalize_table(table)
    for row in table.rows[1:]:
        step_number = _extract_step_number(row.cells[0].text.strip())
        if step_number in STEP_COLORS:
            fill, text = STEP_COLORS[step_number]
            for column_index in (0, 1):
                _shade_cell(row.cells[column_index], fill)
                for run in row.cells[column_index].paragraphs[0].runs:
                    run.font.color.rgb = RGBColor.from_string(text)
                    run.font.bold = True


def _finalize_process_card(table) -> None:
    _finalize_table(table, header_rows=0, font_size=12)
    for row in table.rows:
        _shade_cell(row.cells[0], LIGHT_GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)


def _extract_step_number(text: str) -> int | None:
    digits = "".join(char for char in text if char.isdigit())
    return int(digits) if digits else None


def _iter_document_text(doc: Document) -> Iterable[str]:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    yield text

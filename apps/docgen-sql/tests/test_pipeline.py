from pathlib import Path

from docx import Document

from src.config import load_project_config
from src.run_pipeline import run


def test_pipeline_generates_json_docx_and_audit():
    result = run("input/sql/01_finacle_saldos_decrypt.sql", "config/proyecto.yml")

    assert Path(result["json"]).exists()
    assert Path(result["docx"]).exists()
    assert Path(result["audit"]).exists()
    assert result["audit_passed"] is True

    doc = Document(result["docx"])
    doc_text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert "1. Referencia rapida y ficha del producto" in doc_text
    assert "8. Convenciones y navegacion" in doc_text
    assert "Equipo de datos de Captacion" in doc_text
    assert "Disponible antes de las 08:00 hrs" in doc_text
    assert "paso_02_detalle_cuenta_alnova" in doc_text
    assert "Detalle cuenta ALNOVA" in doc_text
    assert "MAPA DE MODULOS Y SECUENCIA FUNCIONAL DEL PIPELINE" in doc_text
    assert "PRINCIPAL / Principal" in doc_text


def test_step_pipeline_generates_modular_docx():
    result = run(
        None,
        "config/proyecto.yml",
        mode="step",
        module_id="paso_02_detalle_cuenta_alnova",
    )

    assert Path(result["json"]).exists()
    assert Path(result["docx"]).exists()
    assert Path(result["audit"]).exists()
    assert result["mode"] == "step"
    assert result["module"] == "paso_02_detalle_cuenta_alnova"
    assert result["audit_passed"] is True

    doc = Document(result["docx"])
    doc_text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert "Detalle cuenta ALNOVA" in doc_text
    assert "Criterio de aceptacion" in doc_text


def test_step_pipeline_documents_multiple_publications():
    result = run(
        None,
        "config/proyecto.yml",
        mode="step",
        module_id="paso_04_saldo_disponible",
    )

    assert result["audit_passed"] is True
    doc = Document(result["docx"])
    doc_text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    assert "PUBLICACIONES DEL MODULO" in doc_text
    assert "ws_ec_cu_baz_bdclientes.cu_cap_saldos_disponibles_ctas_prev" in doc_text
    assert "ws_ec_cu_baz_bdclientes.cu_cap_saldos_disponibles_ctas" in doc_text
    assert "Intermedia" in doc_text
    assert "Final + compute stats" in doc_text or "Final" in doc_text


def test_step_pipeline_generates_all_modules():
    config = load_project_config("config/proyecto.yml")

    result = run(
        None,
        "config/proyecto.yml",
        mode="step",
        all_steps=True,
    )

    assert result["mode"] == "step"
    assert result["scope"] == "all_steps"
    assert result["modules_requested"] == len(config.flujo.modulos)
    assert result["modules_generated"] + result["modules_failed"] == len(config.flujo.modulos)
    assert len(result["results"]) == result["modules_generated"]
    assert len(result["errors"]) == result["modules_failed"]
    assert result["modules_generated"] >= 1
    assert all(Path(item["json"]).exists() for item in result["results"])
    assert all(Path(item["docx"]).exists() for item in result["results"])
    assert all(Path(item["audit"]).exists() for item in result["results"])
    assert all("module" in item and "sql" in item for item in result["errors"])
